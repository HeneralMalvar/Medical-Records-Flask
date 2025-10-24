from flask import Flask, request, jsonify, render_template, send_file
import sqlite3
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt
from datetime import datetime
import re

app = Flask(__name__)
DB_FILE = "clinic.db"
CSV_FILE = "clinic.csv"

# ---------- DB Helpers ----------
def get_conn():
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_conn()
    c = conn.cursor()

    # Patients table
    c.execute("""
        CREATE TABLE IF NOT EXISTS patients (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            sex TEXT
        )
    """)

    # Visits table (✅ Added remarks column)
    c.execute("""
        CREATE TABLE IF NOT EXISTS visits (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_id INTEGER,
            visit_date TEXT DEFAULT (datetime('now','localtime')),
            age INTEGER,
            address TEXT,
            status TEXT,
            history TEXT,
            pe TEXT,
            diagnosis TEXT,
            management TEXT,
            remarks TEXT,
            FOREIGN KEY (patient_id) REFERENCES patients (id)
        )
    """)

    conn.commit()
    conn.close()
    export_to_csv()

def query_all(query, args=()):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(query, args)
    rows = cur.fetchall()
    conn.close()
    return rows

def execute(query, args=()):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(query, args)
    conn.commit()
    lastrowid = cur.lastrowid
    conn.close()
    export_to_csv()
    return lastrowid

# ---------- CSV Export ----------
def export_to_csv():
    conn = get_conn()
    df = pd.read_sql_query("""
        SELECT p.id as patient_id, p.name, p.sex,
               v.id as visit_id, v.visit_date, v.age as visit_age,
               v.address, v.status, v.history, v.pe, v.diagnosis, v.management, v.remarks
        FROM patients p
        LEFT JOIN visits v ON p.id = v.patient_id
        ORDER BY p.name COLLATE NOCASE ASC, datetime(v.visit_date) DESC
    """, conn)
    conn.close()
    df.to_csv(CSV_FILE, index=False, encoding="utf-8")

init_db()

# ---------- Helper Functions ----------
def format_date(dt_str):
    """Convert stored date to Month Day, Year format."""
    if not dt_str:
        return datetime.now().strftime("%B %d, %Y")
    s = str(dt_str).split(" ")[0]
    try:
        d = datetime.strptime(s, "%Y-%m-%d")
        return d.strftime("%B %d, %Y")
    except Exception:
        return dt_str or datetime.now().strftime("%B %d, %Y")

def replace_underscore_groups_in_paragraph(paragraph, replacements):
    """Replace underscore groups in a paragraph with the given replacements, styled Cambria 14."""
    text = paragraph.text or ""
    parts = re.split(r'(_+)', text)
    result_parts = []
    repl_idx = 0
    for part in parts:
        if re.fullmatch(r'_+', part):
            if repl_idx < len(replacements):
                val = str(replacements[repl_idx] or "")
                result_parts.append(val)
                repl_idx += 1
            else:
                result_parts.append(part)
        else:
            result_parts.append(part)
    new_text = "".join(result_parts)

    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(new_text)
    run.font.name = "Cambria"
    run.font.size = Pt(14)

# ---------- Routes ----------
@app.route("/")
def index():
    return render_template("index.html")

# ---- Patients ----
@app.route("/api/patients", methods=["POST"])
def create_patient():
    data = request.json or {}
    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "Name is required"}), 400
    sex = data.get("sex") or None
    pid = execute("INSERT INTO patients (name, sex) VALUES (?, ?)", (name, sex))
    return jsonify({"message": "Patient created", "patient_id": pid}), 201

@app.route("/api/patients", methods=["GET"])
def get_patients():
    rows = query_all("SELECT id, name, sex FROM patients ORDER BY name COLLATE NOCASE ASC")
    return jsonify([dict(r) for r in rows])

@app.route("/api/patients/<int:pid>", methods=["GET"])
def get_patient(pid):
    rows = query_all("SELECT id, name, sex FROM patients WHERE id=?", (pid,))
    if not rows:
        return jsonify({"error": "Not found"}), 404
    return jsonify(dict(rows[0]))

@app.route("/api/patients/<int:pid>", methods=["PUT"])
def update_patient(pid):
    data = request.json or {}
    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "Name is required"}), 400
    sex = data.get("sex") or None
    execute("UPDATE patients SET name=?, sex=? WHERE id=?", (name, sex, pid))
    return jsonify({"message": "Patient updated"})

@app.route("/api/patients/<int:pid>", methods=["DELETE"])
def delete_patient(pid):
    execute("DELETE FROM visits WHERE patient_id=?", (pid,))
    execute("DELETE FROM patients WHERE id=?", (pid,))
    return jsonify({"message": "Patient and their visits deleted"})

@app.route("/api/patients/search", methods=["GET"])
def search_patients():
    q = (request.args.get("q") or "").strip()
    rows = query_all("SELECT id, name, sex FROM patients WHERE name LIKE ? ORDER BY name COLLATE NOCASE ASC", (f"%{q}%",))
    return jsonify([dict(r) for r in rows])

# ---- Visits ----
@app.route("/api/patients/<int:pid>/visits", methods=["POST"])
def add_visit(pid):
    p = query_all("SELECT id FROM patients WHERE id=?", (pid,))
    if not p:
        return jsonify({"error": "Patient not found"}), 404
    data = request.json or {}
    visit_date = data.get("visit_date")
    age = data.get("age")
    address = data.get("address")
    status = data.get("status")
    history = data.get("history") or ""
    pe = data.get("pe") or ""
    diagnosis = data.get("diagnosis") or ""
    management = data.get("management") or ""
    remarks = data.get("remarks") or ""  # ✅ NEW FIELD

    if visit_date:
        vid = execute("""
            INSERT INTO visits (patient_id, visit_date, age, address, status, history, pe, diagnosis, management, remarks)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (pid, visit_date, age, address, status, history, pe, diagnosis, management, remarks))
    else:
        vid = execute("""
            INSERT INTO visits (patient_id, age, address, status, history, pe, diagnosis, management, remarks)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (pid, age, address, status, history, pe, diagnosis, management, remarks))
    return jsonify({"message": "Visit added", "visit_id": vid})

@app.route("/api/patients/<int:pid>/visits", methods=["GET"])
def get_visits(pid):
    p = query_all("SELECT id, name FROM patients WHERE id=?", (pid,))
    if not p:
        return jsonify({"error": "Patient not found"}), 404
    rows = query_all("""
        SELECT id, visit_date, age, address, status, history, pe, diagnosis, management, remarks
        FROM visits WHERE patient_id=? ORDER BY datetime(visit_date) DESC
    """, (pid,))
    return jsonify([dict(r) for r in rows])

@app.route("/api/visits/<int:vid>", methods=["GET"])
def get_visit(vid):
    rows = query_all("""
        SELECT id, patient_id, visit_date, age, address, status, history, pe, diagnosis, management, remarks
        FROM visits WHERE id=?
    """, (vid,))
    if not rows:
        return jsonify({"error": "Visit not found"}), 404
    return jsonify(dict(rows[0]))

# ---- Update Visit ----
@app.route("/api/visits/<int:vid>", methods=["PUT"])
def update_visit(vid):
    rows = query_all("SELECT id FROM visits WHERE id=?", (vid,))
    if not rows:
        return jsonify({"error": "Visit not found"}), 404

    data = request.json or {}
    visit_date = data.get("visit_date")
    age = data.get("age")
    address = data.get("address") or ""
    status = data.get("status") or ""
    history = data.get("history") or ""
    pe = data.get("pe") or ""
    diagnosis = data.get("diagnosis") or ""
    management = data.get("management") or ""
    remarks = data.get("remarks") or ""  # ✅ NEW FIELD

    try:
        execute("""
            UPDATE visits
            SET visit_date=?, age=?, address=?, status=?, history=?, pe=?, diagnosis=?, management=?, remarks=?
            WHERE id=?
        """, (visit_date, age, address, status, history, pe, diagnosis, management, remarks, vid))
        return jsonify({"message": "Visit updated"})
    except Exception as e:
        app.logger.error(f"Error updating visit {vid}: {e}")
        return jsonify({"error": "Error updating visit"}), 500

# ---- Delete Visit ----
@app.route("/api/visits/<int:vid>", methods=["DELETE"])
def delete_visit(vid):
    rows = query_all("SELECT id FROM visits WHERE id=?", (vid,))
    if not rows:
        return jsonify({"error": "Visit not found"}), 404
    try:
        execute("DELETE FROM visits WHERE id=?", (vid,))
        return jsonify({"message": "Visit deleted"})
    except Exception as e:
        app.logger.error(f"Error deleting visit {vid}: {e}")
        return jsonify({"error": "Error deleting visit"}), 500

# ---- Print Medical Certificate ----
@app.route("/api/visits/<int:vid>/print", methods=["GET"])
def print_medcert(vid):
    """Generate MEDCERT WORK.docx filled with date, name, address, history, diagnosis, and remarks."""
    rows = query_all("""
        SELECT v.*, p.name, p.sex
        FROM visits v
        JOIN patients p ON v.patient_id = p.id
        WHERE v.id=?
    """, (vid,))
    if not rows:
        return jsonify({"error": "Visit not found"}), 404

    data = dict(rows[0])
    name = (data.get("name") or "").strip()
    sex = (data.get("sex") or "").strip()
    age = str(data.get("age") or "").strip()
    address = (data.get("address") or "").strip()
    history = (data.get("history") or "").strip()
    diagnosis = (data.get("diagnosis") or "").strip()
    remarks = (data.get("remarks") or "").strip()  # ✅ NEW FIELD
    visit_date_raw = data.get("visit_date")
    date_text = format_date(visit_date_raw)

    name_with_info = f"{name} {age}/{sex}" if sex or age else name

    doc = Document("MEDCERT WORK.docx")
    processed = set()

    for idx, p in enumerate(doc.paragraphs):
        txt = p.text or ""
        low = txt.lower()

        # --- DATE ---
        if "date" in low and re.search(r'_+', txt):
            replace_underscore_groups_in_paragraph(p, [date_text])
            processed.add(idx)
            continue

        # --- NAME + ADDRESS ---
        if ("certify that" in low or "this is to certify" in low) and re.search(r'_+', txt):
            replace_underscore_groups_in_paragraph(p, [name_with_info, address])
            processed.add(idx)
            continue

        # --- EXAMINED DUE TO ---
        if "examined" in low and "due to" in low and re.search(r'_+', txt):
            replace_underscore_groups_in_paragraph(p, [" " + history])
            processed.add(idx)
            continue

        # --- IMPRESSION ---
        if low.strip().startswith("impression"):
            for j in range(idx + 1, min(idx + 6, len(doc.paragraphs))):
                if j in processed:
                    continue
                p2 = doc.paragraphs[j]
                if re.search(r'_+', p2.text or ""):
                    replace_underscore_groups_in_paragraph(p2, [diagnosis])
                    processed.add(j)
                    break
            continue

        # --- REMARKS --- ✅ NEW LOGIC
        if low.strip().startswith("remarks"):
            for j in range(idx + 1, min(idx + 6, len(doc.paragraphs))):
                if j in processed:
                    continue
                p2 = doc.paragraphs[j]
                if re.search(r'_+', p2.text or ""):
                    replace_underscore_groups_in_paragraph(p2, [remarks])
                    processed.add(j)
                    break
            continue

    # Remove trailing empty paragraphs
    while len(doc.paragraphs) > 0 and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    filename = f"Medical_Certificate_{name.replace(' ', '_')}.docx"

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ---------- Main ----------
if __name__ == "__main__":
    app.config['JSONIFY_PRETTYPRINT_REGULAR'] = True
    app.run(debug=True)
